"""
Document Processing Service

Handles extraction of text and metadata from various document types:
- PDFs (text-based and scanned with OCR)
- Images (JPG, PNG, etc.) with OCR
- Word documents
- Excel spreadsheets
"""

import os
import io
from typing import Dict, Any, List, Optional, BinaryIO
from datetime import datetime
import mimetypes

from PIL import Image
import PyPDF2
from docx import Document
import openpyxl

# OCR imports (optional - will work without if tesseract not installed)
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR not available. Install pytesseract and pdf2image for image/scanned PDF support")


class DocumentProcessor:
    """
    Service for processing various document types and extracting text content
    """

    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'image/jpeg': 'image',
        'image/jpg': 'image',
        'image/png': 'image',
        'image/gif': 'image',
        'image/bmp': 'image',
        'image/tiff': 'image',
        'application/msword': 'doc',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.ms-excel': 'xls',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'text/plain': 'text',
    }

    def __init__(self):
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.ocr_available = OCR_AVAILABLE

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process a document and extract text content and metadata

        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type (optional, will be detected if not provided)

        Returns:
            Dictionary with extracted text, metadata, and processing info
        """
        # Validate file size
        if len(file_content) > self.max_file_size:
            raise ValueError(f"File size exceeds maximum allowed size of {self.max_file_size / 1024 / 1024}MB")

        # Detect content type if not provided
        if not content_type:
            content_type, _ = mimetypes.guess_type(filename)

        # Determine document type
        doc_type = self.SUPPORTED_MIME_TYPES.get(content_type, 'unknown')

        if doc_type == 'unknown':
            raise ValueError(f"Unsupported file type: {content_type}")

        # Process based on document type
        result = {
            "filename": filename,
            "file_size": len(file_content),
            "content_type": content_type,
            "document_type": doc_type,
            "processed_at": datetime.utcnow().isoformat(),
            "text_content": "",
            "metadata": {},
            "pages": 0,
            "extraction_method": "unknown"
        }

        try:
            if doc_type == 'pdf':
                extracted = await self._process_pdf(file_content)
            elif doc_type == 'image':
                extracted = await self._process_image(file_content)
            elif doc_type == 'docx':
                extracted = await self._process_docx(file_content)
            elif doc_type == 'xlsx':
                extracted = await self._process_xlsx(file_content)
            elif doc_type == 'text':
                extracted = await self._process_text(file_content)
            else:
                raise ValueError(f"Processing not implemented for type: {doc_type}")

            result.update(extracted)
            result["success"] = True

        except Exception as e:
            result["success"] = False
            result["error"] = str(e)

        return result

    async def _process_pdf(self, content: bytes) -> Dict[str, Any]:
        """Extract text from PDF (with OCR fallback for scanned PDFs)"""
        result = {
            "text_content": "",
            "pages": 0,
            "metadata": {},
            "extraction_method": "pdf_text"
        }

        try:
            # Try text extraction first
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            result["pages"] = len(pdf_reader.pages)
            result["metadata"] = pdf_reader.metadata or {}

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")

            extracted_text = "\n\n".join(text_parts)

            # If we got very little text, try OCR
            if len(extracted_text.strip()) < 100 and self.ocr_available:
                print(f"PDF appears to be scanned (only {len(extracted_text)} chars), trying OCR...")
                ocr_result = await self._process_pdf_with_ocr(content)
                if ocr_result["text_content"]:
                    result["text_content"] = ocr_result["text_content"]
                    result["extraction_method"] = "pdf_ocr"
                else:
                    result["text_content"] = extracted_text
            else:
                result["text_content"] = extracted_text

        except Exception as e:
            # If text extraction fails, try OCR as fallback
            if self.ocr_available:
                print(f"PDF text extraction failed, trying OCR: {e}")
                ocr_result = await self._process_pdf_with_ocr(content)
                result.update(ocr_result)
            else:
                raise Exception(f"Failed to extract text from PDF: {e}")

        return result

    async def _process_pdf_with_ocr(self, content: bytes) -> Dict[str, Any]:
        """Process PDF using OCR (for scanned documents)"""
        if not self.ocr_available:
            return {"text_content": "", "extraction_method": "ocr_unavailable"}

        try:
            # Convert PDF pages to images
            images = convert_from_bytes(content)

            # OCR each page
            text_parts = []
            for page_num, image in enumerate(images, 1):
                page_text = pytesseract.image_to_string(image)
                if page_text.strip():
                    text_parts.append(f"--- Page {page_num} (OCR) ---\n{page_text}")

            return {
                "text_content": "\n\n".join(text_parts),
                "pages": len(images),
                "extraction_method": "pdf_ocr"
            }

        except Exception as e:
            print(f"OCR processing failed: {e}")
            return {"text_content": "", "extraction_method": "ocr_failed"}

    async def _process_image(self, content: bytes) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        if not self.ocr_available:
            raise Exception("OCR is not available. Install pytesseract and tesseract-ocr")

        try:
            image = Image.open(io.BytesIO(content))

            # Get image metadata
            metadata = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
            }

            # Extract text using OCR
            text = pytesseract.image_to_string(image)

            return {
                "text_content": text,
                "pages": 1,
                "metadata": metadata,
                "extraction_method": "image_ocr"
            }

        except Exception as e:
            raise Exception(f"Failed to process image: {e}")

    async def _process_docx(self, content: bytes) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = Document(io.BytesIO(content))

            # Extract all paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)

            return {
                "text_content": "\n\n".join(text_parts),
                "pages": 1,
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables)
                },
                "extraction_method": "docx_text"
            }

        except Exception as e:
            raise Exception(f"Failed to process DOCX: {e}")

    async def _process_xlsx(self, content: bytes) -> Dict[str, Any]:
        """Extract text from Excel spreadsheet"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)

            text_parts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"--- Sheet: {sheet_name} ---")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)

            return {
                "text_content": "\n".join(text_parts),
                "pages": len(workbook.sheetnames),
                "metadata": {
                    "sheets": workbook.sheetnames
                },
                "extraction_method": "xlsx_text"
            }

        except Exception as e:
            raise Exception(f"Failed to process XLSX: {e}")

    async def _process_text(self, content: bytes) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            text = content.decode('utf-8')

            return {
                "text_content": text,
                "pages": 1,
                "metadata": {},
                "extraction_method": "text_plain"
            }

        except Exception as e:
            raise Exception(f"Failed to process text file: {e}")

    def get_supported_types(self) -> List[str]:
        """Get list of supported file types"""
        return list(self.SUPPORTED_MIME_TYPES.keys())
